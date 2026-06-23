# Генератор документов Word
import os
from docx import Document
from docx.shared import Pt
def create_office_document(content_text: str, filename: str) -> str:
    try:
        doc = Document()
        p = doc.add_paragraph().add_run(content_text)
        p.font.name = 'Times New Roman'; p.font.size = Pt(14)
        path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "Notes", f"{filename}.docx")
        doc.save(path)
        return f"Документ сохранен: Notes/{filename}.docx"
    except Exception as e: return str(e)
